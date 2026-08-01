"""Document generation service: template rendering, LLM drafting, export."""

from __future__ import annotations

import logging
import uuid
from pathlib import Path

from jinja2 import BaseLoader, Environment, StrictUndefined
from markdown import markdown as md_to_html

from app.core.config import get_settings
from app.services.llm import LlmMessage, LlmRole, get_provider_for

logger = logging.getLogger(__name__)

_EXPORT_DIR = Path(get_settings().local_storage_root) / "exports"

_GENERATION_SYSTEM = (
    "你是一個專業文件撰寫助理，服務於仁愛堂社會服務部"
    "（Yan Oi Tong Social Services Division）。\n"
    "請根據使用者提供的指示和參考資料，生成一份正式文件。\n\n"
    "要求：\n"
    "- 使用繁體中文（Traditional Chinese），專有名詞可保留英文\n"
    "- 使用正式公文語氣，結構嚴謹\n"
    "- 如提供了範本，請遵循範本結構\n"
    "- 只返回文件內容，不要加入多餘說明\n"
    "- 文件格式使用 Markdown"
)

_REVISE_SYSTEM = """你是一個專業文件編輯助理。請根據使用者提供的修改指示，對文件進行修訂。
只返回修訂後的文件全文，使用繁體中文。"""


#  ---------------------------------------------------------------------------
#  Template rendering
#  ---------------------------------------------------------------------------
_jinja_env = Environment(
    loader=BaseLoader(),
    undefined=StrictUndefined,
    trim_blocks=True,
    lstrip_blocks=True,
)


def render_template(template_content: str, variables: dict[str, str]) -> str:
    """Render a jinja2 template with the given variables."""
    tmpl = _jinja_env.from_string(template_content)
    return tmpl.render(**variables)


#  ---------------------------------------------------------------------------
#  Generation
#  ---------------------------------------------------------------------------


async def generate_document(
    prompt: str,
    *,
    template_content: str | None = None,
    rendered_template: str | None = None,
    kb_context: str = "",
    model: str = "deepseek-v4-flash",
    temperature: float = 0.3,
    max_tokens: int = 8192,
) -> tuple[str, dict]:
    """Generate a document via LLM, optionally guided by a template."""
    user_parts = []

    if template_content:
        user_parts.append(f"## 文件範本\n\n{template_content}")

    if kb_context:
        user_parts.append(f"## 參考資料\n\n{kb_context}")

    if rendered_template:
        user_parts.append(f"## 已填入變數的範本\n\n{rendered_template}")

    user_parts.append(f"## 生成指示\n\n{prompt}")

    messages = [
        LlmMessage(role=LlmRole.SYSTEM, content=_GENERATION_SYSTEM),
        LlmMessage(role=LlmRole.USER, content="\n\n".join(user_parts)),
    ]

    provider, info = get_provider_for(model)
    resp = await provider.chat(
        messages,
        model=info.id,
        temperature=temperature,
        max_tokens=max_tokens,
    )
    return resp.content, {
        "prompt_tokens": resp.usage.prompt_tokens if resp.usage else 0,
        "completion_tokens": resp.usage.completion_tokens if resp.usage else 0,
        "model": resp.model,
    }


async def revise_document(
    current_content: str,
    revision_prompt: str,
    *,
    model: str = "deepseek-v4-flash",
    temperature: float = 0.3,
    max_tokens: int = 8192,
) -> tuple[str, dict]:
    """Revise an existing document based on user instructions."""
    messages = [
        LlmMessage(role=LlmRole.SYSTEM, content=_REVISE_SYSTEM),
        LlmMessage(
            role=LlmRole.USER,
            content=f"## 原始文件\n\n{current_content}\n\n## 修改指示\n\n{revision_prompt}",
        ),
    ]
    provider, info = get_provider_for(model)
    resp = await provider.chat(
        messages,
        model=info.id,
        temperature=temperature,
        max_tokens=max_tokens,
    )
    return resp.content, {
        "prompt_tokens": resp.usage.prompt_tokens if resp.usage else 0,
        "completion_tokens": resp.usage.completion_tokens if resp.usage else 0,
        "model": resp.model,
    }


#  ---------------------------------------------------------------------------
#  Export
#  ---------------------------------------------------------------------------

_DOC_STYLE = """
<style>
  @page { size: A4; margin: 2.5cm 2cm 2.5cm 2cm; }
  body { font-family: "Noto Sans TC", "Microsoft JhengHei", "PMingLiU", serif;
         font-size: 12pt; line-height: 1.8; color: #1a1a1a; }
  h1 { font-size: 18pt; text-align: center; margin-bottom: 1.5em; }
  h2 { font-size: 14pt; margin-top: 1.5em; }
  h3 { font-size: 12pt; margin-top: 1.2em; }
  table { border-collapse: collapse; width: 100%; margin: 1em 0; }
  th, td { border: 1px solid #333; padding: 8px; text-align: left; }
  th { background-color: #f0f0f0; }
</style>
"""


async def _md_to_docx(markdown_content: str, output_path: Path) -> str:
    """Convert markdown to a .docx file using python-docx."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "PMingLiU"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.8

    for line in markdown_content.split("\n"):
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.font.size = Pt(18)
            run.bold = True
        elif line.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.font.size = Pt(14)
            run.bold = True
        elif line.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.font.size = Pt(12)
            run.bold = True
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("|"):
            continue  # tables are complex; skip for now, use markdown in body
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    _EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = str(output_path)
    doc.save(output_path)
    return output_path


async def _md_to_pdf(markdown_content: str, output_path: Path) -> str:
    """Convert markdown to PDF via HTML + WeasyPrint."""
    html_body = md_to_html(
        markdown_content,
        extensions=["tables", "fenced_code", "codehilite"],
    )
    html = f"""<!DOCTYPE html>
<html lang="zh-Hant">
<head><meta charset="utf-8">{_DOC_STYLE}</head>
<body>{html_body}</body>
</html>"""

    from weasyprint import HTML

    _EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = str(output_path)
    HTML(string=html).write_pdf(output_path)
    return output_path


async def export_document(
    doc_id: uuid.UUID,
    content: str,
    fmt: str,  # "docx" | "pdf" | "both"
) -> dict[str, str | None]:
    """Export a document to the requested format(s). Returns paths keyed by format."""
    result: dict[str, str | None] = {"docx": None, "pdf": None}
    stem = f"{doc_id}"

    if fmt in ("docx", "both"):
        docx_path = _EXPORT_DIR / f"{stem}.docx"
        result["docx"] = await _md_to_docx(content, docx_path)

    if fmt in ("pdf", "both"):
        pdf_path = _EXPORT_DIR / f"{stem}.pdf"
        result["pdf"] = await _md_to_pdf(content, pdf_path)

    return result
